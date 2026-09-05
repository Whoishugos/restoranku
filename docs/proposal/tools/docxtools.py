"""Helpers for surgical edits on the proposal .docx while preserving formatting."""

import copy

from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class Doc:
    def __init__(self, document):
        self.doc = document
        self.body = document.element.body
        self.changes = []

    # ---------- lookup ----------

    def paras(self):
        return self.doc.paragraphs

    def p(self, idx):
        return self.doc.paragraphs[idx]

    def find_para(self, needle, start=0):
        """Index of first paragraph containing needle."""
        for i, p in enumerate(self.doc.paragraphs[start:], start=start):
            if needle in p.text:
                return i
        raise KeyError(f"paragraph containing {needle!r} not found")

    def tbl(self, idx):
        return self.doc.tables[idx]

    # ---------- text edits ----------

    def set_text(self, idx, text, label=None):
        """Replace a paragraph's whole text, keeping the first run's formatting."""
        p = self.p(idx)
        _set_para_text(p, text)
        self.changes.append(label or f"set P{idx:04d}")
        return p

    def sub(self, idx, old, new, label=None, required=True):
        """Substring replacement inside a paragraph, run-aware."""
        p = self.p(idx)
        if old not in p.text:
            if required:
                raise ValueError(f"P{idx:04d}: {old!r} not found in {p.text[:120]!r}")
            return False
        _sub_para(p, old, new)
        self.changes.append(label or f"sub P{idx:04d}: {old[:40]!r}")
        return True

    def append_text(self, idx, extra, label=None):
        p = self.p(idx)
        run = p.runs[-1] if p.runs else None
        if run is None:
            _set_para_text(p, p.text + extra)
        else:
            run.text = run.text + extra
        self.changes.append(label or f"append P{idx:04d}")
        return p

    # ---------- structural edits ----------

    def insert_paras_after(self, idx, items, label=None):
        """items: list of (template_para_index, text). Inserted in order after idx."""
        anchor = self.p(idx)._p
        new_els = []
        for tmpl_idx, text in items:
            el = copy.deepcopy(self.p(tmpl_idx)._p)
            _strip_ids(el)
            newp = Paragraph(el, self.p(idx)._parent)
            _set_para_text(newp, text)
            anchor.addnext(el)
            anchor = el
            new_els.append(el)
        self.changes.append(label or f"insert {len(items)} paras after P{idx:04d}")
        return new_els

    def insert_table_after_para(self, idx, template, widths, rows,
                                header=True, label=None):
        tbl = _build_table(template, widths, rows, header)
        self.p(idx)._p.addnext(tbl)
        self.changes.append(label or "insert table")
        return tbl

    def replace_table(self, table, widths, rows, header=True, label=None):
        old = table._tbl
        new = _build_table(table, widths, rows, header)
        old.addnext(new)
        old.getparent().remove(old)
        self.changes.append(label or "replace table")
        return new

    def drop_table(self, table, label=None):
        el = table._tbl
        el.getparent().remove(el)
        self.changes.append(label or "drop table")

    def clone_para_before(self, src_idx, dst_idx, label=None):
        """Copy a paragraph (including any drawing) and insert it before dst_idx."""
        clone = copy.deepcopy(self.p(src_idx)._p)
        _strip_ids(clone)
        _refresh_drawing_ids(clone)
        self.p(dst_idx)._p.addprevious(clone)
        self.changes.append(label or f"clone P{src_idx:04d} before P{dst_idx:04d}")
        return clone

    def drop_blanks_after(self, idx, max_n=20, label=None):
        """Remove consecutive empty paragraphs (no text, no image) following idx."""
        removed = 0
        while removed < max_n:
            nxt = self.p(idx + 1)
            if nxt.text.strip() or "w:drawing" in nxt._p.xml:
                break
            nxt._p.getparent().remove(nxt._p)
            removed += 1
        self.changes.append(label or f"drop {removed} blank paras after P{idx:04d}")
        return removed

    def drop_para(self, idx, label=None):
        el = self.p(idx)._p
        el.getparent().remove(el)
        self.changes.append(label or f"drop P{idx:04d}")

    # ---------- table cell edits ----------

    def cell(self, table, row, col, text, label=None):
        tc = table.rows[row].cells[col]._tc
        _set_cell_text(tc, text)
        self.changes.append(label or f"cell R{row:02d}C{col}")

    def add_rows(self, table, rows, label=None):
        tmpl = table.rows[-1]._tr
        for data in rows:
            tr = copy.deepcopy(tmpl)
            _strip_ids(tr)
            tcs = tr.findall(qn("w:tc"))
            for tc, val in zip(tcs, data):
                _set_cell_text(tc, val)
            table._tbl.append(tr)
        self.changes.append(label or f"add {len(rows)} rows")


# ---------------------------------------------------------------- internals


_DRAWING_ID = [90000]


def _refresh_drawing_ids(el):
    """Give cloned drawings fresh non-visual ids so Word does not see duplicates."""
    for tag in ("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr",
                "{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr",
                "{http://schemas.openxmlformats.org/drawingml/2006/main}cNvPr"):
        for node in el.iter(tag):
            _DRAWING_ID[0] += 1
            node.set("id", str(_DRAWING_ID[0]))
            if node.get("name"):
                node.set("name", f"image-{_DRAWING_ID[0]}")


def _strip_ids(el):
    for node in el.iter():
        for attr in list(node.attrib):
            if attr.endswith("paraId") or attr.endswith("textId"):
                del node.attrib[attr]


def _set_para_text(p, text):
    runs = p.runs
    if not runs:
        run = p.add_run(text)
        return run
    runs[0].text = text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)
    return runs[0]


def _sub_para(p, old, new):
    """Replace old->new across run boundaries, keeping the first affected run's style."""
    runs = p.runs
    joined = "".join(r.text for r in runs)
    pos = joined.find(old)
    if pos < 0:
        raise ValueError("substring vanished")
    end = pos + len(old)
    # Map char offsets to runs.
    offsets = []
    acc = 0
    for r in runs:
        offsets.append((acc, acc + len(r.text), r))
        acc += len(r.text)
    first = None
    for start, stop, r in offsets:
        if stop <= pos or start >= end:
            continue
        head = r.text[: max(0, pos - start)]
        tail = r.text[max(0, end - start):] if end - start < len(r.text) else ""
        if first is None:
            first = r
            r.text = head + new + tail
        else:
            r.text = head + tail
    return p


def _set_cell_text(tc, text):
    """Set cell content; '\n' splits into multiple paragraphs."""
    ps = tc.findall(qn("w:p"))
    tmpl = copy.deepcopy(ps[0])
    _strip_ids(tmpl)
    for p in ps:
        tc.remove(p)
    lines = str(text).split("\n")
    for line in lines:
        el = copy.deepcopy(tmpl)
        _strip_ids(el)
        tc.append(el)
        # reuse the first run of the cloned paragraph
        runs = el.findall(qn("w:r"))
        if not runs:
            continue
        for extra in runs[1:]:
            el.remove(extra)
        t = runs[0].find(qn("w:t"))
        if t is None:
            t = runs[0].makeelement(qn("w:t"), {})
            runs[0].append(t)
        t.text = line
        t.set(
            "{http://www.w3.org/XML/1998/namespace}space", "preserve"
        )
    return tc


def _build_table(template_table, widths, rows, header=True):
    """Build a w:tbl cloning the template's look, with given column widths + data."""
    tmpl = template_table._tbl
    tbl = tmpl.makeelement(qn("w:tbl"), {})

    tblPr = copy.deepcopy(tmpl.find(qn("w:tblPr")))
    w = tblPr.find(qn("w:tblW"))
    if w is not None:
        w.set(qn("w:w"), str(float(sum(widths))))
    tbl.append(tblPr)

    grid = tmpl.makeelement(qn("w:tblGrid"), {})
    for width in widths:
        gc = tmpl.makeelement(qn("w:gridCol"), {})
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    tbl.append(grid)

    tmpl_rows = tmpl.findall(qn("w:tr"))
    head_tr = tmpl_rows[0]
    body_tr = tmpl_rows[1] if len(tmpl_rows) > 1 else tmpl_rows[0]
    head_tc = head_tr.findall(qn("w:tc"))[0]
    body_tcs = body_tr.findall(qn("w:tc"))
    body_first = body_tcs[0]
    body_rest = body_tcs[1] if len(body_tcs) > 1 else body_tcs[0]

    for ri, data in enumerate(rows):
        is_head = header and ri == 0
        src_tr = head_tr if is_head else body_tr
        tr = tmpl.makeelement(qn("w:tr"), {})
        trPr = src_tr.find(qn("w:trPr"))
        if trPr is not None:
            newTrPr = copy.deepcopy(trPr)
            cs = newTrPr.find(qn("w:cantSplit"))
            if cs is not None:
                cs.set(qn("w:val"), "1")
            if is_head:
                th = newTrPr.find(qn("w:tblHeader"))
                if th is not None:
                    th.set(qn("w:val"), "1")
            tr.append(newTrPr)
        for ci, val in enumerate(data):
            if is_head:
                src_tc = head_tc
            else:
                src_tc = body_first if ci == 0 else body_rest
            tc = copy.deepcopy(src_tc)
            _strip_ids(tc)
            _set_cell_text(tc, val)
            tr.append(tc)
        tbl.append(tr)
    return tbl
